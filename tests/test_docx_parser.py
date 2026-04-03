from pathlib import Path

from docx import Document

from app.engine.masker import MaskingEngine
from app.parsers.docx_parser import process_docx


def test_docx_cross_run_masking(tmp_path: Path):
    source = tmp_path / 'sample.docx'
    output = tmp_path / 'sample_masked.docx'

    doc = Document()
    p = doc.add_paragraph()
    p.add_run('张')
    p.add_run('三')
    p.add_run('的手机号是13812345678')
    doc.save(source)

    engine = MaskingEngine(Path(__file__).resolve().parents[1] / 'config')
    process_docx(source, output, engine)

    masked = Document(output)
    text = '\n'.join(p.text for p in masked.paragraphs)
    assert '张三' not in text
    assert '138****5678' in text
