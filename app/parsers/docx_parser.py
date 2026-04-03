from pathlib import Path
from typing import List, Tuple

from docx import Document

from app.engine.masker import MaskingEngine
from app.models.report import HitDetail


def _mask_paragraph_runs(para, engine: MaskingEngine, location_prefix: str) -> List[HitDetail]:
    details: List[HitDetail] = []
    full_text = ''.join(run.text for run in para.runs)
    if not full_text.strip():
        return details

    masked_text, hits = engine.desensitize_text(full_text, location=location_prefix)
    if not hits:
        return details

    if para.runs:
        para.runs[0].text = masked_text
        for run in para.runs[1:]:
            run.text = ''
    details.extend(hits)
    return details


def process_docx(file_path: Path, output_path: Path, engine: MaskingEngine) -> Tuple[Path, List[HitDetail]]:
    doc = Document(file_path)
    details: List[HitDetail] = []

    for p_idx, para in enumerate(doc.paragraphs, start=1):
        details.extend(_mask_paragraph_runs(para, engine, f"Paragraph_{p_idx}"))

    for t_idx, table in enumerate(doc.tables, start=1):
        for row_idx, row in enumerate(table.rows, start=1):
            for col_idx, cell in enumerate(row.cells, start=1):
                for p_idx, para in enumerate(cell.paragraphs, start=1):
                    details.extend(_mask_paragraph_runs(para, engine, f"Table_{t_idx}_R{row_idx}C{col_idx}_P{p_idx}"))

    doc.save(output_path)
    return output_path, details
